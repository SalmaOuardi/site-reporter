"""Generate DOCX reports from template using extracted fields."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def _strip_placeholder(text: str, field: str) -> str:
    """Remove placeholder tokens for a given field from cell text."""

    return text.replace(f"{{{field}}}", "").replace(f"[{field}]", "").strip()


def _replace_generic_placeholders(text: str, fields: Dict[str, str], field_mapping: Dict[str, str]) -> str:
    """Fallback replacement for any loose placeholders outside the main table layout."""

    for template_field, extracted_field in field_mapping.items():
        value = fields.get(extracted_field, "")
        text = text.replace(f"{{{template_field}}}", value)
        text = text.replace(f"[{template_field}]", value)
    return text


def _normalize_label(text: str) -> str:
    """Normalize a label for matching (remove colon/spacing, lowercase)."""

    return text.replace(":", "").strip().lower()


def _set_cell_shading(cell, fill: str) -> None:
    """Apply a background fill color to a cell."""

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _style_header_cell(cell) -> None:
    """Style section header cells for clarity."""

    _set_cell_shading(cell, "e8efff")
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(4)
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(12, 44, 92)
            run.font.size = Pt(12)


def _style_label_cell(cell) -> None:
    """Style label cells (left column)."""

    _set_cell_shading(cell, "f4f6fb")
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(32, 45, 64)


def _style_value_cell(cell) -> None:
    """Style value cells (right column)."""

    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            run.font.bold = False
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(30, 38, 50)


def generate_incident_docx(fields: Dict[str, str], template_path: Path) -> bytes:
    """Fill the incident Word template with extracted data and return bytes."""

    doc = Document(template_path)

    # Field mapping: template field name -> extracted field name
    field_mapping = {
        "Nom du chantier": "Nom du chantier",
        "Nom de l'incident": "Nom de l'incident",
        "Emetteur du signalement": "Emetteur du signalement",
        "Date de découverte": "Date de découverte",
        "Heure de découverte": "Heure de découverte",
        "Adresse": "Adresse",
        "Nature de l'incident": "Nature de l'incident",
        "Description de l'incident": "Description de l'incident",
        "Risques identifiés": "Risques identifiés",
        "Actions à réaliser": "Actions à réaliser",
        "Niveau d'urgence": "Niveau d'urgence",
        "Personnes prévenues": "Personnes prévenues",
    }

    # Replace placeholders in tables, ensuring values land in column 2 (right-hand cell)
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            matched = False

            if len(cells) >= 2:
                left_cell, right_cell = cells[0], cells[1]
                left_label = _normalize_label(left_cell.text)

                for template_field, extracted_field in field_mapping.items():
                    norm_field = template_field.lower()
                    value = fields.get(extracted_field, "")

                    # Match by label presence in the left column or explicit placeholder
                    if norm_field in left_label or f"{{{template_field}}}" in left_cell.text or f"[{template_field}]" in left_cell.text:
                        left_cell.text = _strip_placeholder(left_cell.text, template_field)
                        right_cell.text = value
                        matched = True
                        break

                    # If the placeholder lives in the right cell, swap it for the value
                    if f"{{{template_field}}}" in right_cell.text or f"[{template_field}]" in right_cell.text:
                        left_cell.text = _strip_placeholder(left_cell.text, template_field)
                        right_cell.text = _strip_placeholder(right_cell.text, template_field)
                        right_cell.text = value
                        matched = True
                        break

            # Fallback: replace any straggling placeholders (single-column rows, titles, etc.)
            if not matched:
                for cell in cells:
                    cell.text = _replace_generic_placeholders(cell.text, fields, field_mapping)

    # Presentation: style headers and columns to keep things readable
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) == 1:
                _style_header_cell(cells[0])
            else:
                _style_label_cell(cells[0])
                for cell in cells[1:]:
                    _style_value_cell(cell)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.read()
